"""
ExiusCart Company Letterhead
Layout: Logo + ExiusCart name side-by-side left | company contact right
Style: clean white, thin purple accent line, like FAIRAM reference
"""

import io
from PIL import Image, ImageDraw
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PURPLE  = RGBColor(0x6B, 0x3F, 0xD9)
DARK    = RGBColor(0x12, 0x12, 0x1A)
GREY    = RGBColor(0x6B, 0x72, 0x80)
LT_GREY = RGBColor(0xCC, 0xCC, 0xCC)


# ── Real logo from SVG polygon data ──────────────────────────────────────────

def make_logo_png(size=220) -> bytes:
    S = size / 100.0
    def sc(pts):
        return [(x * S, y * S) for x, y in pts]

    img = Image.new("RGBA", (size, size), (0, 0, 0, 0))
    d   = ImageDraw.Draw(img)

    d.polygon(sc([(50,4),(88,27),(88,73),(50,96),(12,73),(12,27)]),
              fill=(0x6B, 0x3F, 0xD9, 255))
    d.polygon(sc([(50,28),(69,39),(50,50),(31,39)]),
              fill=(255, 255, 255, 255))
    d.polygon(sc([(69,39),(69,62),(50,73),(50,50)]),
              fill=(255, 255, 255, 191))
    d.polygon(sc([(31,39),(31,62),(50,73),(50,50)]),
              fill=(255, 255, 255, 127))

    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf.read()


# ── Helpers ───────────────────────────────────────────────────────────────────

def no_borders(table):
    tbl  = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tb = OxmlElement("w:tblBorders")
    for side in ("top","left","bottom","right","insideH","insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "none")
        el.set(qn("w:sz"),    "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tb.append(el)
    tblPr.append(tb)


def paragraph_border(para, side, color, sz=6):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    el   = OxmlElement(f"w:{side}")
    el.set(qn("w:val"),   "single")
    el.set(qn("w:sz"),    str(sz))
    el.set(qn("w:space"), "1")
    el.set(qn("w:color"), color)
    pBdr.append(el)
    pPr.append(pBdr)


def add_run(para, text, size=10, color=None, bold=False, italic=False):
    r = para.add_run(text)
    r.font.size  = Pt(size)
    r.font.bold  = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    return r


# ── Main ──────────────────────────────────────────────────────────────────────

def build():
    doc     = Document()
    section = doc.sections[0]
    section.page_width    = Cm(21)
    section.page_height   = Cm(29.7)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)
    section.top_margin    = Cm(1.2)
    section.bottom_margin = Cm(1.8)

    logo_bytes = make_logo_png(220)

    # ════════════════════════════════════════════════
    # HEADER
    # Layout (same as FAIRAM reference):
    #   [LOGO | ExiusCart name + tagline]   [contact block right-aligned]
    # ════════════════════════════════════════════════
    hdr = section.header
    for p in hdr.paragraphs:
        p.clear()

    # Outer 2-col table: left brand | right contact
    ht = hdr.add_table(rows=1, cols=2, width=Cm(16.6))
    ht.autofit = False
    no_borders(ht)
    ht.columns[0].width = Cm(10.5)
    ht.columns[1].width = Cm(6.1)

    # ── LEFT SIDE: logo tightly beside text ──
    lc = ht.cell(0, 0)
    lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Inner table: [logo cell] [name cell] — no gap between them
    inner = lc.add_table(rows=1, cols=2)
    inner.autofit = False
    no_borders(inner)
    inner.columns[0].width = Cm(1.5)  # logo
    inner.columns[1].width = Cm(9.0)  # name

    # Logo cell
    logo_c = inner.cell(0, 0)
    logo_c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    lp = logo_c.paragraphs[0]
    lp.paragraph_format.space_before = Pt(0)
    lp.paragraph_format.space_after  = Pt(0)
    lp.add_run().add_picture(io.BytesIO(logo_bytes), width=Cm(1.25))

    # Name cell — "ExiusCart" bold + tagline
    name_c = inner.cell(0, 1)
    name_c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    np1 = name_c.paragraphs[0]
    np1.paragraph_format.space_before = Pt(0)
    np1.paragraph_format.space_after  = Pt(1)
    np1.paragraph_format.left_indent  = Cm(0.25)
    r_e = np1.add_run("Exius")
    r_e.bold = True
    r_e.font.size = Pt(24)
    r_e.font.color.rgb = PURPLE
    r_c = np1.add_run("Cart")
    r_c.bold = True
    r_c.font.size = Pt(24)
    r_c.font.color.rgb = DARK

    np2 = name_c.add_paragraph()
    np2.paragraph_format.space_before = Pt(0)
    np2.paragraph_format.space_after  = Pt(0)
    np2.paragraph_format.left_indent  = Cm(0.25)
    rt = np2.add_run("All-in-one business management platform")
    rt.font.size  = Pt(8)
    rt.font.color.rgb = GREY

    # ── RIGHT SIDE: company contact block ──
    rc = ht.cell(0, 1)
    rc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    contacts = [
        ("www.exiuscart.com",          False),
        ("support@exiuscart.com",      False),
        ("Fairam Private Limited",     True),
        ("Dubai, United Arab Emirates",False),
    ]
    for i, (text, bold) in enumerate(contacts):
        p = rc.paragraphs[0] if i == 0 else rc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        r = p.add_run(text)
        r.font.size = Pt(8.5)
        r.font.bold = bold
        r.font.color.rgb = DARK if bold else GREY

    # ── Purple accent line ──
    accent = hdr.add_paragraph()
    accent.paragraph_format.space_before = Pt(8)
    accent.paragraph_format.space_after  = Pt(0)
    paragraph_border(accent, "bottom", "6B3FD9", sz=10)

    # ════════════════════════════════════════════════
    # BODY
    # ════════════════════════════════════════════════

    def blank(n=1, before=0, after=0):
        for _ in range(n):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(before)
            p.paragraph_format.space_after  = Pt(after)

    def line(text="", size=11, color=None, bold=False, before=0, after=4):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after  = Pt(after)
        if text:
            r = p.add_run(text)
            r.font.size = Pt(size)
            r.bold = bold
            if color:
                r.font.color.rgb = color
        return p

    def field_line(label, before=0, after=3):
        """Label: ____________"""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after  = Pt(after)
        rl = p.add_run(label)
        rl.font.size  = Pt(10.5)
        rl.font.color.rgb = DARK
        rv = p.add_run("  " + "_" * 46)
        rv.font.size  = Pt(10.5)
        rv.font.color.rgb = LT_GREY
        return p

    # Date + Ref row
    blank(1, before=12)
    dr = doc.add_paragraph()
    dr.paragraph_format.space_before = Pt(0)
    dr.paragraph_format.space_after  = Pt(6)
    add_run(dr, "Date:",  10.5, DARK,    bold=True)
    add_run(dr, "  " + "_"*28,  10.5, LT_GREY)
    add_run(dr, "       Ref No:", 10.5, DARK, bold=True)
    add_run(dr, "  " + "_"*20,  10.5, LT_GREY)

    # To block
    to_p = doc.add_paragraph()
    to_p.paragraph_format.space_before = Pt(10)
    to_p.paragraph_format.space_after  = Pt(2)
    add_run(to_p, "To,", 11, DARK, bold=True)

    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2)
        add_run(p, "_" * 52, 11, LT_GREY)

    # Subject
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_before = Pt(12)
    sub_p.paragraph_format.space_after  = Pt(8)
    add_run(sub_p, "Subject: ", 11, DARK, bold=True)
    add_run(sub_p, "_" * 58,    11, LT_GREY)

    # Light divider under subject
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(0)
    sep.paragraph_format.space_after  = Pt(8)
    paragraph_border(sep, "bottom", "E5E7EB", sz=4)

    # Greeting
    line("Dear Sir / Madam,", 11, DARK, before=4, after=10)

    # Body writing lines
    for _ in range(15):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(5)
        add_run(p, "_" * 88, 11, LT_GREY)

    # Closing
    line("Yours sincerely,", 11, DARK, before=20, after=2)
    blank(4)

    # Signature line
    sig_p = doc.add_paragraph()
    sig_p.paragraph_format.space_before = Pt(0)
    sig_p.paragraph_format.space_after  = Pt(3)
    paragraph_border(sig_p, "bottom", "6B3FD9", sz=4)
    add_run(sig_p, " " * 48)

    line("Authorised Signatory", 10, GREY, before=2, after=1)

    ep = doc.add_paragraph()
    ep.paragraph_format.space_before = Pt(0)
    ep.paragraph_format.space_after  = Pt(0)
    add_run(ep, "ExiusCart", 10, PURPLE, bold=True)
    add_run(ep, "  ·  Fairam Private Limited", 10, DARK)

    # ════════════════════════════════════════════════
    # FOOTER
    # ════════════════════════════════════════════════
    ftr = section.footer
    for p in ftr.paragraphs:
        p.clear()

    # Top purple line
    ftop = ftr.paragraphs[0]
    ftop.paragraph_format.space_before = Pt(0)
    ftop.paragraph_format.space_after  = Pt(5)
    paragraph_border(ftop, "top", "6B3FD9", sz=6)

    # Brand line
    f1 = ftr.add_paragraph()
    f1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f1.paragraph_format.space_before = Pt(0)
    f1.paragraph_format.space_after  = Pt(1)
    add_run(f1, "ExiusCart",             9, PURPLE, bold=True)
    add_run(f1, "  —  a brand of  ",     9, DARK)
    add_run(f1, "Fairam Private Limited",9, DARK,   bold=True)

    # Contact line
    f2 = ftr.add_paragraph()
    f2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f2.paragraph_format.space_before = Pt(0)
    f2.paragraph_format.space_after  = Pt(1)
    add_run(f2, "www.exiuscart.com  ·  support@exiuscart.com  ·  Developed by NexCodeNova",
            7.5, GREY)

    # Copyright
    f3 = ftr.add_paragraph()
    f3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f3.paragraph_format.space_before = Pt(0)
    f3.paragraph_format.space_after  = Pt(0)
    add_run(f3, "(c) 2026 Fairam Private Limited. All rights reserved.  |  Confidential",
            7, LT_GREY)

    # Default font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    out = r"m:\Projects for NexCode Nova\ExiusCart\ExiusCart_Letterhead.docx"
    doc.save(out)
    print(f"DONE -> {out}")


if __name__ == "__main__":
    build()
